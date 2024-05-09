import os
from init_bot import bot
import docx
from docx import Document
from docxtpl import DocxTemplate
import sqlite3 as sl
import time
from progress.bar import IncrementalBar
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import threading
from secrets import db


class Thread(threading.Thread):
    def __init__(self, t, *args):
        threading.Thread.__init__(self, target=t, args=args)
        self.start()
count = 0
lock = threading.Lock()


def month_plan(call):
    try:
        chat_id = call.chat.id
    except:
        chat_id = call.message.chat.id
    
    month = ""
    month_str = ""
    year = ""
    
    with bot.retrieve_data(user_id=call.from_user.id, chat_id=chat_id) as data:
        year = str(data.get("looking_year", None))
        month = str(data.get("looking_month", None))
        month_str = data.get("looking_month_str", None)

    
    if int(month) < datetime.now().month and int(year) <= datetime.now().year:
        bot.send_message(chat_id, "❗️ Мероприятия прошедших месяцев можно посмотреть в архиве", parse_mode='html')
        return
    
    lock.acquire()
    
    export_msg = bot.send_message(chat_id, "Выгружаю план...\nКак будет готово, я сразу пришлю файл.", parse_mode='html')
    bot.add_data(call.from_user.id, chat_id, export_msg=export_msg.id, userid=call.from_user.id)
    
    date = "%" + month + "." + year + "%"

    con = sl.connect(db)
    cursor = con.cursor()
    sql = """
    SELECT 
    day_of_week,
    SUBSTR(date, 1, INSTR(date, '.') - 1) || '.' || SUBSTR(date, INSTR(date, '.') + 1, INSTR(SUBSTR(date, INSTR(date, '.') + 1), '.') - 1) AS date,
    start_time || CHAR(10) || ' - ' || CHAR(10) || end_time AS time,
    title,
    count_students,
    count_visitors,
    place,
    seniors
    FROM events
    WHERE date LIKE ?
    ORDER BY date, start_time"""
    data = (date, )
    cursor.execute(sql, data)
    events = cursor.fetchall()
    
    if len(events) == 0:
        with bot.retrieve_data(user_id=call.from_user.id, chat_id=chat_id) as data:
            export_message = data.get("export_msg", None)
            bot.delete_message(chat_id, export_message)
        bot.send_message(chat_id, "Мероприятия отсутствуют", parse_mode='html')
        lock.release()
        return

    # bar = IncrementalBar('Countdown', max = len(events))

    doc = Document('blanc_events_plan.docx')

    for paragraph in doc.paragraphs:
        if '{{month}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{month}}', f'{month_str}')
            run = paragraph.runs[0]
            run.bold = True
        if '{{year}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{year}}', f'{year}')
            run = paragraph.runs[0]
            run.bold = True
            
    table0 = doc.tables[0]

    for row in table0.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                if '{{year}}' in run.text:
                    run.text = run.text.replace('{{year}}', f'{year}')

    table1 = doc.tables[1]

    for index, row in enumerate(events):
        new_row = table1.add_row()
        for index, cell in enumerate(new_row.cells):
            cell.text = f"{row[index]}"
        # bar.next()
        
    for row_index in range(2, len(table1.rows)):
        for cell_index in range(8):
            current_cell = table1.rows[row_index].cells[cell_index]
            prevous_cell = table1.rows[row_index - 1].cells[cell_index]
            for run in current_cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            if cell_index == 0 or cell_index == 1:
                current_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in current_cell.paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                if current_cell.text in prevous_cell.text and cell_index < 2:
                    current_cell.text = ""
                    merged_cell = table1.cell(row_index, cell_index).merge(table1.cell(row_index - 1, cell_index))
                    # cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell_index == 2:
                current_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if current_cell.text == f"\n - \n":
                    current_cell.text = ''
            if cell_index == 4 or cell_index == 5:
                # cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                current_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'month_{month}_{call.message.message_id}.docx')

    # bar.finish()
    
    time.sleep(1)
    
    try:
        with bot.retrieve_data(user_id=call.from_user.id, chat_id=chat_id) as data:
            export_message = data.get("export_msg", None)
            bot.delete_message(chat_id, export_message)
    except:
        pass
    
    bot.send_document(chat_id, open(rf'month_{month}_{call.message.message_id}.docx', 'rb'))
    
    os.remove(f'month_{month}_{call.message.message_id}.docx')
    
    lock.release()
